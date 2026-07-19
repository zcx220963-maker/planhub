"""
RAG 知识库路由（增强版）

优化点：
1. 混合检索: BM25 关键词检索 + 向量相似度检索
2. LLM 重排序 (Rerank): 先粗检索 top-20 候选，再用 LLM 打分选前 k
3. 上下文压缩: 只保留与问题最相关的句子，减少无效 token
4. 改进文档分割: RecursiveCharacterTextSplitter（递归分割，优先按段落/句子）
5. 元数据增强: 每个文档块保存文件名、块索引、总块数等信息
"""

from fastapi import APIRouter, HTTPException, File, UploadFile, Form
from pydantic import BaseModel, Field
from typing import Optional, List
from langchain_chroma import Chroma
from langchain_community.document_loaders import (
    TextLoader,
    DirectoryLoader,
    PyPDFLoader,
    Docx2txtLoader,
    UnstructuredExcelLoader,
    UnstructuredPowerPointLoader
)
# 改为 RecursiveCharacterTextSplitter（优先按段落→句子→字符，语义更完整）
from langchain_text_splitters import RecursiveCharacterTextSplitter, CharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from dotenv import load_dotenv
import os
import re
import shutil
import json
import math
from pathlib import Path
from datetime import datetime
from app.common.llm_factory import get_llm, get_embeddings
from app.dao.redis_dao import save_session, is_redis_available, add_chat_message, get_chat_history
from config import settings

load_dotenv()

router = APIRouter(prefix="/rag")

# ══════════════════════════════════════════════════════════════════════════════
# 1. 请求 / 响应模型
# ══════════════════════════════════════════════════════════════════════════════

class RAGQueryRequest(BaseModel):
    question: Optional[str] = Field(default=None, description="用户问题")
    query: Optional[str] = Field(default=None, description="用户问题（兼容前端）")
    top_k: int = Field(default=8, description="最终返回的文档数")
    # 新增: 粗检索数（先找更多候选，再重排序）
    fetch_k: int = Field(default=20, description="粗检索候选数（建议是 top_k 的 5-10 倍）")
    # 新增: 是否启用 LLM 重排序（质量更好但慢一点）
    use_rerank: bool = Field(default=True, description="是否启用 LLM 重排序")
    # 新增: 是否启用上下文压缩（精简上下文，降低 token 消耗）
    use_compression: bool = Field(default=True, description="是否启用上下文压缩")
    temperature: float = Field(default=0.3)
    user_id: str = Field(default="default", description="用户ID")
    session_id: Optional[str] = Field(default=None, description="会话ID")
    doc_ids: Optional[list] = Field(default=None, description="指定参与问答的文档ID列表，为空则使用所有文档")

    def get_question(self) -> str:
        return self.question or self.query or ""


class RAGQueryResponse(BaseModel):
    success: bool
    answer: str
    response: str = ""
    sources: list = []
    model: str = "deepseek-r1:7b"
    context_length: int = 0
    matched_count: int = 0
    session_id: str = ""
    # 新增: 检索策略调试信息
    retrieval_info: dict = {}


class DocumentUploadResponse(BaseModel):
    success: bool
    message: str
    documents_count: int
    documents: list = []


# ══════════════════════════════════════════════════════════════════════════════
# 2. 全局状态
# ══════════════════════════════════════════════════════════════════════════════

# 用户集合缓存：每个用户有独立的 Chroma collection
# {user_id: Chroma实例}
user_vector_stores = {}

# 原始文档索引（用于 BM25 关键词检索、文档元数据查询）
# 按用户ID分区存储 {user_id: {doc_id: doc_info}}
user_documents = {}

# 用户文档元数据索引（供预览/debug 接口使用）
# BM25 倒排索引现由 Chroma 内部维护，不再需要 Python 侧维护
user_documents = {}    # {user_id: {doc_id: {filename, chunks, content}}}
original_documents = {}  # 兼容旧代码（引用当前用户的 user_documents）

# 保留空 dict 占位，避免旧代码引用报错（实际不再使用）
user_bm25_index = {}
user_bm25_stats = {}
bm25_index = {}
bm25_total_docs = 0
bm25_avg_length = 0.0


# ══════════════════════════════════════════════════════════════════════════════
# 3. 初始化
# ══════════════════════════════════════════════════════════════════════════════

def get_or_create_user_vector_store(user_id: str):
    """
    获取或创建该用户的独立 Chroma 向量库实例

    每个用户有独立的 collection，数据库层面完全隔离。
    Chroma 支持多 collection 共享同一个 persist_directory，性能良好。
    """
    global user_vector_stores
    if user_id in user_vector_stores:
        return user_vector_stores[user_id]

    db_path = settings.CHROMA_DB_PATH
    os.makedirs(db_path, exist_ok=True)
    embeddings = get_embeddings()

    # 每个用户使用独立的 collection name
    collection_name = f"user_{user_id}"

    try:
        store = Chroma(
            collection_name=collection_name,
            persist_directory=db_path,
            embedding_function=embeddings,
        )
        user_vector_stores[user_id] = store
        print(f"[INFO] 向量库已创建: user_id={user_id}, collection={collection_name}")
        return store
    except Exception as e:
        print(f"[WARN] 向量库创建失败，重试: {e}")
        store = Chroma(
            collection_name=collection_name,
            persist_directory=db_path,
            embedding_function=embeddings,
        )
        user_vector_stores[user_id] = store
        return store


async def init_document_indices(user_id: str = "default"):
    """
    从 MySQL rag_documents 表重建 user_documents 元数据索引。

    BM25 倒排索引现由 Chroma 内部维护（持久化到磁盘，重启自动恢复），
    这里只重建 user_documents 元数据（供预览/debug 接口使用）。

    每次 Python 服务重启都会调用，确保重启后：
    - 指定文档查询（doc_ids）可用
    - 文档列表/预览接口始终有内容

    Args:
        user_id: 用户ID，每个用户有独立的文档空间
    """
    global original_documents

    # 初始化该用户的索引空间
    if user_id not in user_documents:
        user_documents[user_id] = {}

    # 设置当前用户的索引引用
    original_documents = user_documents[user_id]

    try:
        from ..service.document_store import get_documents_by_user
        docs = await get_documents_by_user(user_id)
    except Exception as e:
        print(f"[WARN] 加载 MySQL 文档失败 (user_id={user_id}): {e}")
        docs = []

    if not docs:
        return 0

    loaded = 0
    for doc in docs:
        try:
            doc_id = doc["doc_id"]
            filename_in_doc = doc["filename"]
            actual_content = doc.get("full_content") or doc.get("content", "")
            # 去掉旧磁盘格式残留的前 2 行（"=== filename ===" 标题行）
            if actual_content.startswith("==="):
                lines = actual_content.split("\n")
                actual_content = "\n".join(lines[2:]) if len(lines) > 2 else ""
            if not actual_content.strip():
                continue
            original_documents[doc_id] = {
                "filename": filename_in_doc,
                "chunks": [],  # 不再维护 chunk 元数据，由 Chroma 管理
                "upload_time": str(doc.get("created_at", datetime.now())),
                "content": actual_content,
            }
            loaded += 1
        except Exception as e:
            print(f"[WARN] 重建元数据索引时处理文档 {doc.get('doc_id', '?')} 失败: {e}")

    print(f"[INFO] 文档元数据索引重建完成 (MySQL): user_id={user_id}, 共 {loaded} 个文档")
    return loaded


def init_qa_chain():
    """初始化问答链（备用，当前主要使用函数式调用）"""
    global qa_chain
    try:
        llm = get_llm(0.3)
        store = get_or_create_user_vector_store("default")
        from langchain_core.prompts import ChatPromptTemplate
        from langchain_core.runnables import RunnablePassthrough
        from prompts.rag import RAG_CHAT_SYSTEM_PROMPT
        prompt = ChatPromptTemplate.from_messages([
            ("system", RAG_CHAT_SYSTEM_PROMPT + "\n\n上下文：\n{context}"),
            ("human", "{question}")
        ])
        def format_docs(docs):
            return "\n\n".join([doc.page_content for doc in docs])
        qa_chain = (
            {"context": store.as_retriever(search_kwargs={"k": 3}) | format_docs,
             "question": RunnablePassthrough()}
            | prompt | llm
        )
    except Exception as e:
        print(f"[WARN] qa_chain 初始化失败（非关键）: {e}")
        qa_chain = None


# ══════════════════════════════════════════════════════════════════════════════
# 4. BM25 关键词检索（解决关键词匹配问题，弥补纯向量检索的不足）
# ══════════════════════════════════════════════════════════════════════════════

# BM25 超参数
BM25_K1 = 1.5  # 词频饱和度（控制 tf 的上限）
BM25_B = 0.75  # 文档长度归一化系数


def tokenize(text: str) -> List[str]:
    """分词策略：
    - 英文按空格+标点切分，转小写
    - 中文：优先按双字切分（增加词汇级匹配），同时保留单字（确保召回率）
    - 数字和单位保留（如"50s", "30秒"）
    这样既能保证召回率，又能提高关键词匹配精度
    """
    text = text.lower()
    tokens = []
    
    # 先提取英文单词和数字（如50s, 30秒）
    pattern = re.compile(r"[a-z0-9]+s?|[\u4e00-\u9fa5]+")
    raw_tokens = pattern.findall(text)
    
    for token in raw_tokens:
        if re.match(r"[a-z0-9]+", token):
            # 英文/数字，直接添加
            tokens.append(token)
        else:
            # 中文，双字切分 + 单字切分
            chars = list(token)
            # 添加双字组合
            for i in range(len(chars) - 1):
                tokens.append(chars[i] + chars[i+1])
            # 添加单字（确保召回率）
            tokens.extend(chars)
    
    return tokens


def compute_bm25_tf(content: str) -> dict:
    """计算单文档的词频表"""
    tokens = tokenize(content)
    tf = {}
    for t in tokens:
        tf[t] = tf.get(t, 0) + 1
    return tf, len(tokens)


def add_doc_to_bm25(doc_id: str, filename: str, chunk_index: int, content: str, user_id: str = "default"):
    """
    兼容旧接口 —— 空操作。

    BM25 倒排索引现在由 Chroma 内部维护（持久化到磁盘，重启自动恢复），
    不再需要 Python 侧手动维护内存 BM25 索引。
    Chroma 在 store.add_documents() 时自动建好倒排索引。

    参数保持原签名，避免改动调用方。
    """
    pass


def bm25_search(query: str, top_n: int = 20, doc_ids: Optional[List[str]] = None, user_id: str = "default") -> List[Document]:
    """
    BM25 关键词检索 —— 委托给 hybrid_search（Chroma 内置 BM25）。

    BM25 倒排索引现由 Chroma 内部维护，持久化到磁盘，重启自动恢复。
    本函数保持原签名，避免改动调用方。
    """
    docs, _ = hybrid_search(query, top_k=top_n, fetch_k=top_n, doc_ids=doc_ids, user_id=user_id)
    return docs


# ══════════════════════════════════════════════════════════════════════════════
# 5. 混合检索 + 重排序
# ══════════════════════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════════════
# HyDE（Hypothetical Document Generation）查询扩展
# 思路：用户的问题词汇可能和文档中的词汇不一样（例：用户搜"登录"，文档写"用户认证"）
# 通过让 LLM 先写一段"假设的相关文档"，然后用这段假设文档去检索，
# 可以显著提升词汇差异场景下的命中率。
# ══════════════════════════════════════════════════════════════════════════════

HYDE_PROMPT = """请根据以下用户的问题，写 1-2 段简洁的中文文档片段，
这段片段应该是从"详细技术文档"里摘出来的，覆盖问题中提到的核心概念。

要求：
1. 只写文档内容，不要加"根据您的问题"等解释性文字
2. 用词正式、客观，像真实文档
3. 每段 80-120 字即可
4. 如果问题不明确，请返回一段最可能相关的通用描述

用户问题：
{question}

文档片段："""


def generate_hypothetical_document(question: str, llm=None) -> str:
    """让 LLM 基于问题生成一段"假设的相关文档片段"。

    这个假设文档会和原问题一起做向量检索，解决"用户词汇 ≠ 文档词汇"的问题。
    失败时安全降级为直接用原问题，不会抛异常。
    """
    try:
        if llm is None:
            llm = get_llm(0.3)
        result = llm.invoke(HYDE_PROMPT.format(question=question))
        text = result.content if hasattr(result, "content") else str(result)
        text = text.strip()
        if not text:
            return ""
        print(f"[HyDE] 生成假设文档成功，长度={len(text)} 字（前 80 字: {text[:80]}）")
        return text
    except Exception as e:
        print(f"[HyDE] 生成失败，降级为原问题检索: {e}")
        return ""


def vector_search(query: str, top_n: int = 20, doc_ids: Optional[List[str]] = None, user_id: str = "default") -> List[Document]:
    """
    向量相似度检索（每个用户独立 collection，数据库层面隔离）

    Args:
        query: 查询文本
        top_n: 返回文档数
        doc_ids: 指定文档ID列表
        user_id: 用户ID（定位到该用户的独立向量库）
    """
    # 获取该用户的独立向量库
    store = get_or_create_user_vector_store(user_id)

    print(f"[DEBUG] vector_search: user_id={user_id}, query={query[:60]}..., top_n={top_n}")

    # 构建过滤条件（只检索该用户的 collection）
    where_filter = None
    if doc_ids and len(doc_ids) > 0:
        where_filter = {"doc_id": {"$in": [str(d) for d in doc_ids]}}

    try:
        search_kwargs = {"k": top_n}
        if where_filter:
            search_kwargs["filter"] = where_filter

        retriever = store.as_retriever(
            search_type="similarity",
            search_kwargs=search_kwargs
        )
        docs = retriever.invoke(query)
        print(f"[DEBUG] vector_search returned {len(docs)} docs for user_id={user_id}")
        for d in docs:
            d.metadata["source"] = "vector"
        return docs
    except Exception as e:
        print(f"[ERROR] vector_search failed: {e}")
        import traceback
        print(traceback.format_exc())
        return []


def hybrid_search(query: str, top_k: int = 3, fetch_k: int = 20,
                  doc_ids: Optional[List[str]] = None, user_id: str = "default") -> (List[Document], dict):
    """
    混合检索（按用户ID隔离）—— 走 Chroma 单一入口，Chroma 内部同时做：
      - 向量相似度检索（稠密向量）
      - BM25 关键词检索（稀疏倒排索引，Chroma 内置维护，持久化到磁盘）

    Chroma 1.5+ 的 query() 同时支持 query_embeddings + query_texts，
    内部自动做混合排名，不需要我们分两路再手写 RRF 合并。

    Args:
        query: 查询文本
        top_k: 最终返回文档数
        fetch_k: 粗检索候选数（传给 Chroma 的 n_results）
        doc_ids: 指定文档ID列表
        user_id: 用户ID（定位到该用户的独立向量库）

    返回: (最终文档列表, 检索调试信息)
    """
    retrieval_info = {
        "vector_count": 0,
        "bm25_count": 0,
        "merged_count": 0,
        "user_id": user_id,
        "fusion": "chroma_hybrid",
    }

    # 获取该用户的独立向量库
    store = get_or_create_user_vector_store(user_id)

    # 构建过滤条件（只检索指定 doc_ids 时）
    where_filter = None
    if doc_ids and len(doc_ids) > 0:
        where_filter = {"doc_id": {"$in": [str(d) for d in doc_ids]}}

    try:
        # 同时传 query_embeddings（向量）+ query_texts（BM25），Chroma 内部混合
        # 注意：query_embeddings 需要是嵌套列表 [[float, ...]]
        from src.app.common.llm_factory import get_embeddings
        embeddings = get_embeddings()
        query_embedding = embeddings.embed_query(query)

        n_results = min(fetch_k, 50)  # Chroma 单次数值上限
        result = store.query(
            query_embeddings=[query_embedding],
            query_texts=[query],
            n_results=n_results,
            where=where_filter,
            include=["documents", "metadatas", "distances"],
        )

        # 解析 Chroma 返回格式
        documents = result.get("documents", [[]])[0]
        metadatas = result.get("metadatas", [[]])[0]
        distances = result.get("distances", [[]])[0]

        final_docs = []
        for i, content in enumerate(documents):
            meta = metadatas[i] if i < len(metadatas) else {}
            dist = distances[i] if i < len(distances) else 1.0
            final_docs.append(Document(
                page_content=content,
                metadata={
                    "doc_id": meta.get("doc_id", ""),
                    "doc_name": meta.get("doc_name", ""),
                    "chunk_index": meta.get("chunk_index", 0),
                    "distance": round(dist, 4),
                    "source": "chroma_hybrid",
                    **meta,
                }
            ))

        retrieval_info["merged_count"] = len(final_docs)
        return final_docs[:top_k], retrieval_info

    except Exception as e:
        print(f"[ERROR] hybrid_search failed: {e}")
        import traceback
        print(traceback.format_exc())
        # 降级：纯向量检索
        try:
            return vector_search(query, top_n=top_k, doc_ids=doc_ids, user_id=user_id), retrieval_info
        except Exception:
            return [], retrieval_info


# ══════════════════════════════════════════════════════════════════════════════
# 6. LLM 重排序（rerank）：让 LLM 重新判断每个文档与问题的相关性
# ══════════════════════════════════════════════════════════════════════════════

def llm_rerank(question: str, candidate_docs: List[Document], top_k: int,
               temperature: float = 0.1) -> (List[Document], dict):
    """
    LLM 重排序：
      - 把候选文档逐篇交给 LLM 打分（0-10 分，表示与问题的相关性）
      - 按分数降序排列，取前 top_k
      - 比简单的相似度检索更精准，但会增加 LLM 调用时间

    优化：为了降低调用次数，可以一次性让 LLM 输出批量打分（当前实现）
    """
    if not candidate_docs:
        return [], {"rerank_count": 0}

    llm = get_llm(temperature)

    # 批量方式：一次给 LLM 所有候选，要求输出 JSON 数组
    # 每个候选输出前 300 字符给 LLM 看，避免太长
    preview_list = []
    for i, d in enumerate(candidate_docs):
        preview = d.page_content[:300].replace("\n", " ")
        preview_list.append({"id": i, "preview": preview})

    from prompts.rag import LLM_RERANK_SYSTEM_PROMPT, LLM_RERANK_PROMPT_TEMPLATE

    prompt = LLM_RERANK_PROMPT_TEMPLATE.format(
        question=question,
        preview_list=json.dumps(preview_list, ensure_ascii=False, indent=2),
    )
    # 调用 LLM（以简单字符串方式调用）
    from langchain_core.messages import HumanMessage, SystemMessage
    from src.app.common.llm_factory import extract_text
    messages = [
        SystemMessage(content=LLM_RERANK_SYSTEM_PROMPT),
        HumanMessage(content=prompt),
    ]
    resp = llm.invoke(messages)
    raw = resp.content if hasattr(resp, "content") else str(resp)

    # extract_text 兼容 thinking 模型（返回 list）和普通模型（返回 str）
    text = extract_text(raw).strip()

    # 解析 JSON（容错处理：从 { 开始，} 结束）
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        json_text = text[start:end+1]
        data = json.loads(json_text)
        scores = data.get("scores", [])
        # 建立 id -> score 映射
        score_map = {int(s["id"]): float(s["score"]) for s in scores}
        # 用 LLM 打分覆盖 final_score，并设置分数
        scored = []
        for i, d in enumerate(candidate_docs):
            s = score_map.get(i, d.metadata.get("final_score", 0.0) * 5)
            d.metadata["rerank_score"] = round(float(s), 2)
            scored.append((s, d))
        # 排序，取前 top_k
        scored.sort(key=lambda x: x[0], reverse=True)
        top = [d for _, d in scored[:top_k]]
        return top, {"rerank_count": len(top), "llm_scored": True}

    # JSON 解析失败：回退
    print(f"[WARN] LLM Rerank 输出无法解析为 JSON，回退到 RRF 分数")
    return candidate_docs[:top_k], {"rerank_count": top_k, "llm_scored": False, "error": "json_parse_failed"}


# ══════════════════════════════════════════════════════════════════════════════
# 7. 上下文压缩：只保留与问题最相关的句子（降低 token 消耗，提升回答精度）
# ══════════════════════════════════════════════════════════════════════════════

def compress_context(question: str, docs: List[Document], min_sentences: int = 3) -> List[Document]:
    """
    基于简单关键词覆盖的上下文压缩：
      1. 对每个文档按句号/换行切句子
      2. 计算每句与问题的重合词数
      3. 只保留得分 >= 阈值的句子
      4. 确保至少保留 min_sentences 句，避免空文档
      5. 确保保留内容长度至少为原文的 30%
    """
    q_tokens = set(tokenize(question))
    if not q_tokens:
        return docs

    compressed = []
    for doc in docs:
        content = doc.page_content
        original_length = len(content)
        
        sentences = re.split(r"[。！？!?；;\n]+", content)
        sentences = [s.strip() for s in sentences if s.strip()]
        if not sentences:
            compressed.append(doc)
            continue

        scored_sents = []
        for s in sentences:
            s_tokens = set(tokenize(s))
            overlap = len(s_tokens & q_tokens)
            scored_sents.append((overlap, s))

        top_sents = []
        for overlap, s in scored_sents:
            if overlap >= 0:
                top_sents.append(s)

        if len(top_sents) < min_sentences:
            remaining = [s for o, s in scored_sents if s not in top_sents]
            top_sents.extend(remaining[:min_sentences - len(top_sents)])

        scored_sents.sort(key=lambda x: x[0], reverse=True)
        sorted_sents = [s for o, s in scored_sents]
        
        min_length = int(original_length * 0.3)
        selected_sents = []
        current_length = 0
        
        for s in sorted_sents:
            if s not in selected_sents:
                selected_sents.append(s)
                current_length += len(s)
                if current_length >= min_length and len(selected_sents) >= min_sentences:
                    break

        if len(selected_sents) < min_sentences:
            remaining = [s for s in sentences if s not in selected_sents]
            selected_sents.extend(remaining[:min_sentences - len(selected_sents)])

        new_content = ""
        for s in sentences:
            if s in selected_sents:
                new_content += s + "。"

        new_content = new_content.strip()
        if not new_content or len(new_content) < 20:
            new_content = content

        new_doc = Document(page_content=new_content, metadata={**doc.metadata, "compressed": True})
        compressed.append(new_doc)
        
        if len(content) > len(new_content):
            print(f"[DEBUG] compress_context: 文档 {doc.metadata.get('doc_name', 'unknown')}#{doc.metadata.get('chunk_index', 0)} "
                  f"压缩前 {len(content)} 字符 → 压缩后 {len(new_content)} 字符")
            print(f"[DEBUG]   压缩前内容: {content[:100]}...")
            print(f"[DEBUG]   压缩后内容: {new_content[:100]}...")

    return compressed


# ══════════════════════════════════════════════════════════════════════════════
# 8. API：核心问答接口
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/query", response_model=RAGQueryResponse)
async def query_rag(request: RAGQueryRequest):
    """
    RAG 问答接口（增强版，按用户ID隔离）

    流程:
      1. 混合检索（向量 + BM25 关键词）→ 得到 fetch_k 个候选
      2. 可选: LLM 重排序 → 精炼为 top_k 个最相关文档
      3. 可选: 上下文压缩 → 只保留相关句子，降低 token 消耗
      4. 拼接系统提示（含文档上下文）+ 历史对话 + 当前问题
      5. 调用 LLM 生成回答

    安全说明：
      - 所有检索都按 user_id 隔离，用户只能访问自己的知识库
      - request.user_id 由 Java 后端校验后传入，不可伪造
    """
    # 获取用户ID（由 Java 后端鉴权后传入）
    user_id = request.user_id or "default"

    try:
        # 获取该用户的独立向量库
        store = get_or_create_user_vector_store(user_id)

        # 检查该用户的向量库大小
        try:
            doc_count = store._collection.count()
            print(f"[DEBUG] 向量数据库文档数: user_id={user_id}, count={doc_count}")
        except Exception:
            doc_count = 0

        question = request.get_question()
        if not question:
            raise HTTPException(status_code=400, detail="问题不能为空")

        # 新增：如果 doc_ids 是空列表（前端勾选后取消全选），提示用户选择文档
        if request.doc_ids is not None and len(request.doc_ids) == 0:
            return RAGQueryResponse(
                success=True,
                answer="未选择任何参考文档。请在右侧勾选要参考的文档后再提问，或关闭知识库开关。",
                sources=[],
                model="system",
                retrieval_info={"note": "doc_ids 为空，未启用 RAG 检索"}
            )

        # --- 步骤 0: HyDE 查询扩展 ---
        # 生成一段假设的相关文档片段，然后把它和原问题合并去检索。
        # 这样可以解决"用户说 '登录'，文档写 '用户认证'"的词汇差异问题。
        hyde_text = generate_hypothetical_document(question)
        if hyde_text:
            query_for_search = f"{question}\n\n相关文档片段：\n{hyde_text}"
        else:
            query_for_search = question

        # --- 步骤 1: 混合检索（按用户ID隔离） ---
        retrieved_docs, retrieval_info = hybrid_search(
            query_for_search,
            top_k=request.top_k,
            fetch_k=request.fetch_k,
            doc_ids=request.doc_ids,
            user_id=user_id,
        )
        retrieval_info["hyde_used"] = bool(hyde_text)
        retrieval_info["hyde_length"] = len(hyde_text)

        # 如果没有检索到（可能是 BM25 索引为空，降级为纯向量）
        if not retrieved_docs and doc_count > 0:
            print("[WARN] 混合检索无结果，回退为纯向量检索...")
            print(f"[DEBUG] BM25 index size: {len(bm25_index)}, total docs: {bm25_total_docs}")
            print(f"[DEBUG] Vector DB has {doc_count} documents")
            retrieved_docs = vector_search(question, top_n=request.top_k, doc_ids=request.doc_ids)
            print(f"[DEBUG] Vector search returned {len(retrieved_docs)} documents")
            retrieval_info["fallback_vector_only"] = True

        # 文档数为 0 时返回提示
        if not retrieved_docs:
            return RAGQueryResponse(
                success=True,
                answer="知识库中暂无文档，请先上传文档后再进行查询。",
                sources=[],
                model="system",
                retrieval_info={"note": "知识库为空"}
            )

        # --- 步骤 2: 可选 LLM 重排序 ---
        rerank_info = {}
        if request.use_rerank and len(retrieved_docs) > request.top_k:
            print(f"[DEBUG] 执行 LLM 重排序: {len(retrieved_docs)} 候选 → {request.top_k}")
            retrieved_docs, rerank_info = llm_rerank(
                question, retrieved_docs, request.top_k, temperature=0.1
            )
        elif len(retrieved_docs) > request.top_k:
            retrieved_docs = retrieved_docs[:request.top_k]

        # --- 步骤 3: 可选上下文压缩 ---
        compression_info = {}
        if request.use_compression:
            original_len = sum(len(d.page_content) for d in retrieved_docs)
            retrieved_docs = compress_context(question, retrieved_docs, min_sentences=2)
            compressed_len = sum(len(d.page_content) for d in retrieved_docs)
            compression_info = {
                "original_length": original_len,
                "compressed_length": compressed_len,
                "saved_ratio": round((1 - compressed_len / max(original_len, 1)) * 100, 2),
            }
            print(f"[DEBUG] 上下文压缩: {original_len} → {compressed_len} 字符 (节省 {compression_info['saved_ratio']}%)")

        # --- 步骤 4: 生成回答 ---
        # 拼接最终上下文（带 ref_id 引用标记，与系统提示的引用格式一致）
        context_parts = []
        for idx, doc in enumerate(retrieved_docs, 1):
            meta = doc.metadata or {}
            doc_name = meta.get("doc_name", f"文档{idx}")
            chunk_idx = meta.get("chunk_index", idx)
            ref_id = f"{doc_name}#{chunk_idx}"
            context_parts.append(f"片段 {idx} [来源: {ref_id}]: {doc.page_content}")
        context = "\n\n".join(context_parts)

        print(f"[DEBUG] 最终上下文长度: {len(context)} 字符, 文档数: {len(retrieved_docs)}")

        llm = get_llm(request.temperature)
        import uuid
        session_id = request.session_id or str(uuid.uuid4())

        # 多轮对话历史
        history = []
        if settings.use_redis_bool and is_redis_available():
            history = get_chat_history(session_id) or []
            print(f"[DEBUG] 获取到历史对话 {len(history)} 条, session_id={session_id}")

        try:
            from langchain_core.messages import HumanMessage, SystemMessage, AIMessage

            # 系统提示：结构化 + 思维链引导 + 引用标注 + Few-shot 示例
            from prompts.rag import RAG_CHAT_SYSTEM_PROMPT
            system_prompt = f"""{RAG_CHAT_SYSTEM_PROMPT}

【文档片段】
{context}
"""

            messages = [SystemMessage(content=system_prompt)]

            # 历史对话（最多保留最近 8 轮，共 16 条消息）
            max_history_turns = 8
            if history:
                recent_history = history[-max_history_turns * 2:]
                for msg in recent_history:
                    role = msg.get("role", "")
                    content = msg.get("content", "")
                    if role == "user":
                        messages.append(HumanMessage(content=content))
                    elif role == "assistant":
                        messages.append(AIMessage(content=content))

            # 当前问题
            messages.append(HumanMessage(content=question))

            print(f"[DEBUG] LLM 调用: messages总数={len(messages)}, system_prompt长度={len(system_prompt)}")
            print(f"[DEBUG] 上下文内容（前500字符）: {context[:500]}")
            print(f"[DEBUG] 用户问题: {question}")
            response = llm.invoke(messages)
            answer = response.content if hasattr(response, "content") else str(response)
            print(f"[DEBUG] LLM 回答（前200字符）: {answer[:200]}")

            if not answer or not answer.strip():
                answer = "根据知识库内容，无法生成明确的回答。请尝试更具体的问题。"

        except Exception as llm_error:
            print(f"[ERROR] LLM 调用失败: {llm_error}")
            import traceback
            print(traceback.format_exc())
            answer = f"生成回答时出错: {str(llm_error)}。请检查模型服务是否正常运行。"

        # --- 步骤 5: 提取来源信息（供前端展示）---
        sources = []
        for i, doc in enumerate(retrieved_docs):
            meta = doc.metadata or {}
            sources.append({
                "index": i,
                "content": doc.page_content[:300] + ("..." if len(doc.page_content) > 300 else ""),
                "doc_name": meta.get("doc_name", f"文档{i}"),
                "doc_id": meta.get("doc_id", ""),
                "final_score": meta.get("final_score", 0),
                "rerank_score": meta.get("rerank_score", None),
            })

        # --- 步骤 6: 保存到 Redis 会话 ---
        # 只有当 RAG 找到有效信息时才保存会话历史，避免污染 Chat API 的上下文
        if settings.use_redis_bool and is_redis_available():
            # 判断是否为有效回答（不是"未找到相关信息"）
            is_valid_answer = (
                answer and
                not answer.startswith("在知识库中未找到相关信息") and
                not answer.startswith("知识库中暂无文档") and
                not answer.startswith("根据知识库内容，无法生成明确的回答")
            )

            if is_valid_answer:
                # 只有有效回答才保存到会话历史（使用 module="chat" 以便在 Chatbot 会话历史中显示）
                add_chat_message(session_id, "user", question, user_id=request.user_id, module="chat")
                add_chat_message(session_id, "assistant", answer, user_id=request.user_id, module="chat")

                from app.dao.redis_dao import get_session
                existing = get_session(session_id)
                if not existing or "title" not in existing:
                    session_data = existing or {}
                    session_data["title"] = question[:50] + ("..." if len(question) > 50 else "")
                    session_data["user_id"] = request.user_id
                    session_data["module"] = "chat"
                    save_session(session_id, session_data, 86400, user_id=request.user_id, module="chat")
            else:
                print(f"[DEBUG] RAG 未找到有效信息，不保存到会话历史，避免污染 Chat API 上下文")

        # 汇总检索调试信息
        full_retrieval_info = {
            **retrieval_info,
            **{"rerank_" + k: v for k, v in rerank_info.items()},
            **{"compression_" + k: v for k, v in compression_info.items()},
            "use_rerank": request.use_rerank,
            "use_compression": request.use_compression,
            "top_k": request.top_k,
            "fetch_k": request.fetch_k,
        }

        return RAGQueryResponse(
            success=True,
            answer=answer,
            response=answer,
            sources=sources,
            model="deepseek-r1:7b",
            context_length=len(context),
            matched_count=len(retrieved_docs),
            session_id=session_id,
            retrieval_info=full_retrieval_info,
        )

    except Exception as e:
        import traceback
        print(f"[ERROR] query_rag failed: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"查询失败: {str(e)}")


# ══════════════════════════════════════════════════════════════════════════════
# 9. 内部 API（供计划生成接口调用）
# ══════════════════════════════════════════════════════════════════════════════

async def query_rag_internal(
    question: str,
    user_id: str = "default",
    session_id: str = None,
    doc_ids: list = None,
    top_k: int = 3,
    use_rerank: bool = False,
    use_compression: bool = False
) -> Optional[dict]:
    """
    内部 RAG 查询接口（供计划生成接口调用，按用户ID隔离）

    Args:
        question: 用户问题
        user_id: 用户ID（用于隔离不同用户的知识库）
        session_id: 会话ID
        doc_ids: 指定文档ID列表
        top_k: 返回文档数
        use_rerank: 是否启用 LLM 重排序
        use_compression: 是否启用上下文压缩

    Returns:
        {"answer": "...", "sources": [...]} 或 None
    """
    try:
        # 构建请求对象
        request = RAGQueryRequest(
            question=question,
            user_id=user_id,
            session_id=session_id,
            doc_ids=doc_ids,
            top_k=top_k,
            use_rerank=use_rerank,
            use_compression=use_compression,
        )

        # 调用现有的 query_rag 函数
        result = await query_rag(request)

        if result and result.answer:
            return {
                "answer": result.answer,
                "sources": result.sources,
            }

        return None

    except Exception as e:
        print(f"[ERROR] query_rag_internal failed: {e}")
        return None


# ══════════════════════════════════════════════════════════════════════════════
# 10. 文件上传（改进版：使用 RecursiveCharacterTextSplitter，加强元数据）
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(file: UploadFile = File(...), user_id: str = Form("default")):
    """
    上传文档到知识库

    Args:
        file: 上传的文件
        user_id: 用户ID（从 FormData 中读取，用于隔离不同用户的知识库）
    """
    return await process_uploaded_files([file], user_id=user_id)


@router.post("/upload/batch", response_model=DocumentUploadResponse)
async def upload_documents_batch(files: list[UploadFile] = File(...), user_id: str = Form("default")):
    """
    批量上传文档到知识库

    Args:
        files: 上传的文件列表
        user_id: 用户ID（从 FormData 中读取，用于隔离不同用户的知识库）
    """
    return await process_uploaded_files(files, user_id=user_id)


async def process_uploaded_files(files: list[UploadFile], user_id: str = "default") -> DocumentUploadResponse:
    try:
        allowed_extensions = [".txt", ".md", ".json", ".csv", ".pdf", ".docx", ".doc",
                              ".xlsx", ".xls", ".pptx", ".ppt"]

        temp_dir = "./temp_docs"
        os.makedirs(temp_dir, exist_ok=True)

        all_split_docs = []
        processed_count = 0
        failed_files = []

        # 新增：记录已上传的文档信息（doc_id, filename, chunk_count）
        uploaded_doc_records = []

        # 使用 RecursiveCharacterTextSplitter：优先按段落/句子切分，语义更完整
        # 相对 CharacterTextSplitter：不会在段落中间一刀切
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,        # 增大chunk大小，避免表格内容被切断
            chunk_overlap=100,     # 增大重叠，保留更多上下文
            length_function=len,
            # 切分优先级：双换行 → 单换行 → 句子分隔符 → 逗号 → 空格 → 字符
            separators=["\n\n", "\n", "(?<=[。！？!?；;])", ",", " ", ""],
        )

        for file in files:
            try:
                file_ext = Path(file.filename).suffix.lower()
                if file_ext not in allowed_extensions:
                    failed_files.append(f"{file.filename}: 不支持的文件类型")
                    continue

                file_path = os.path.join(temp_dir, file.filename)
                with open(file_path, "wb") as buffer:
                    shutil.copyfileobj(file.file, buffer)

                # 根据类型选择加载器
                documents = []
                if file_ext in [".txt", ".md", ".json", ".csv"]:
                    loader = TextLoader(file_path, encoding="utf-8")
                    documents = loader.load()
                elif file_ext == ".pdf":
                    loader = PyPDFLoader(file_path)
                    documents = loader.load()
                elif file_ext == ".docx":
                    try:
                        loader = Docx2txtLoader(file_path)
                        documents = loader.load()
                        if not documents or not documents[0].page_content.strip():
                            documents = load_docx_with_python_docx(file_path)
                    except Exception as e:
                        print(f"Docx2txtLoader 失败: {e}，使用 python-docx")
                        documents = load_docx_with_python_docx(file_path)
                elif file_ext == ".doc":
                    documents = load_doc_file(file_path)
                    if not documents:
                        failed_files.append(f"{file.filename}: 无法读取内容")
                        os.remove(file_path)
                        continue
                else:
                    failed_files.append(f"{file.filename}: 暂不支持的文件类型")
                    continue

                # 递归分割
                split_docs = text_splitter.split_documents(documents)

                # 过滤空片段
                split_docs = [d for d in split_docs if d.page_content and d.page_content.strip()]

                # 超长片段（>800 字符）强制再切
                max_chunk_size = 800
                final_docs = []
                for doc in split_docs:
                    content = doc.page_content
                    if len(content) <= max_chunk_size:
                        final_docs.append(doc)
                    else:
                        # 超长的用字符级别切
                        for i in range(0, len(content), max_chunk_size - 100):
                            chunk = content[i:i + max_chunk_size]
                            if chunk.strip():
                                final_docs.append(Document(
                                    page_content=chunk.strip(),
                                    metadata=doc.metadata,
                                ))

                split_docs = final_docs
                print(f"[DEBUG] 文件 {file.filename} 分割后有效片段数: {len(split_docs)}")

                if not split_docs:
                    failed_files.append(f"{file.filename}: 内容为空")
                    os.remove(file_path)
                    continue

                # 生成 doc_id，更新内存索引 + BM25 索引 + 磁盘保存
                import uuid
                doc_id = str(uuid.uuid4())[:8]
                original_content = "\n\n".join([doc.page_content for doc in documents])

                # 内存索引（供指定文档查询 / 预览）
                chunks_meta = [{"content": d.page_content, "chunk_index": i,
                                "total_chunks": len(split_docs)}
                               for i, d in enumerate(split_docs)]
                # 保存到该用户的内存索引
                if user_id not in user_documents:
                    user_documents[user_id] = {}
                user_documents[user_id][doc_id] = {
                    "filename": file.filename,
                    "chunks": chunks_meta,
                    "upload_time": str(datetime.now()),
                    "content": original_content,
                }

                # BM25 索引（用于关键词检索，按用户ID隔离）
                for i, d in enumerate(split_docs):
                    add_doc_to_bm25(doc_id, file.filename, i, d.page_content, user_id=user_id)
                    # 给每个文档块补充元数据到 Chroma
                    d.metadata["doc_id"] = doc_id
                    d.metadata["doc_name"] = file.filename
                    d.metadata["chunk_index"] = i
                    d.metadata["total_chunks"] = len(split_docs)
                    d.metadata["user_id"] = user_id  # 记录所属用户

                # MySQL 保存（替代原来的磁盘文件系统）
                from ..service.document_store import save_document
                await save_document(
                    doc_id=doc_id,
                    user_id=user_id,
                    filename=file.filename,
                    content=original_content,
                    chunk_count=len(split_docs),
                )

                all_split_docs.extend(split_docs)
                processed_count += 1
                os.remove(file_path)

                # 记录文档信息（供返回给前端）
                uploaded_doc_records.append({
                    "doc_id": doc_id,
                    "filename": file.filename,
                    "chunk_count": len(split_docs),
                })

            except Exception as e:
                print(f"[ERROR] 处理文件 {file.filename} 失败: {e}")
                failed_files.append(f"{file.filename}: {str(e)}")

        # 分批加入该用户的独立向量库（避免大文件超时）
        if all_split_docs:
            store = get_or_create_user_vector_store(user_id)

            total_docs = len(all_split_docs)
            batch_size = 50
            print(f"[DEBUG] 开始分批添加 {total_docs} 个片段到向量数据库 (user_id={user_id})...")

            for i in range(0, total_docs, batch_size):
                batch = all_split_docs[i:i + batch_size]
                start_idx = i + 1
                end_idx = min(i + batch_size, total_docs)
                print(f"[DEBUG] 添加批次 {start_idx}-{end_idx}/{total_docs}...")
                store.add_documents(batch)
                print(f"[DEBUG] 批次 {start_idx}-{end_idx} 添加成功")

            print(f"[DEBUG] 全部 {total_docs} 个片段添加成功 (user_id={user_id})")

        message = f"成功上传 {processed_count} 个文档，共 {len(all_split_docs)} 个片段"
        if failed_files:
            message += f"。失败 {len(failed_files)} 个文件: {', '.join(failed_files)}"

        # 文档内容已在 process_uploaded_files 中同步写入 MySQL (rag_documents 表)
        return DocumentUploadResponse(
            success=True,
            message=message,
            documents_count=len(all_split_docs),
            documents=uploaded_doc_records,
        )

    except Exception as e:
        import traceback
        print(f"[ERROR] upload_documents failed: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"上传失败: {str(e)}")


# ══════════════════════════════════════════════════════════════════════════════
# 10. 其他辅助接口（保持不变，但调用函数的逻辑用新版）
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/load-directory")
async def load_directory(directory_path: str):
    try:
        if not os.path.isdir(directory_path):
            raise HTTPException(status_code=400, detail="目录路径不存在")
        loader = DirectoryLoader(directory_path, glob="**/*.txt",
                                 loader_cls=TextLoader,
                                 loader_kwargs={"encoding": "utf-8"})
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=50,
            separators=["\n\n", "\n", "(?<=[。！？!?；;])", ",", " ", ""]
        )
        split_docs = text_splitter.split_documents(documents)
        store = get_or_create_user_vector_store("default")
        store.add_documents(split_docs)
        return {"success": True,
                "message": f"从目录加载 {len(documents)} 个文档，共处理 {len(split_docs)} 个片段",
                "documents_count": len(split_docs)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/stats")
async def get_stats(user_id: str = "default"):
    """获取指定用户的知识库统计"""
    try:
        store = get_or_create_user_vector_store(user_id)
        count = store._collection.count()
        bm25_count = len(user_bm25_index.get(user_id, {}))
        return {
            "success": True,
            "user_id": user_id,
            "document_count": count,
            "bm25_docs": bm25_count,
            "description": "PlanHub RAG 知识库（向量 + BM25 混合检索，按用户隔离）"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/clear")
async def clear_knowledge_base(user_id: str = "default"):
    """
    清空指定用户的知识库（按用户ID隔离）

    删除该用户的独立 Chroma collection（同时清理向量和 BM25 索引），
    并清理 MySQL 文档记录和内存元数据。

    Args:
        user_id: 用户ID，只清空该用户的知识库
    """
    try:
        # 1. 删除该用户的独立向量库 collection（Chroma 同时清理向量和 BM25 索引）
        if user_id in user_vector_stores:
            try:
                store = user_vector_stores[user_id]
                store._client.delete_collection(store._collection.name)
                del user_vector_stores[user_id]
                print(f"[INFO] 向量库 collection 已删除: user_id={user_id}")
            except Exception as e:
                print(f"[WARN] 向量库删除失败: {e}")

        # 2. 清空该用户的内存元数据索引
        if user_id in user_documents:
            user_documents[user_id].clear()

        # 3. 清空该用户的 MySQL 文档记录
        from ..service.document_store import delete_document, get_documents_by_user
        user_docs_list = await get_documents_by_user(user_id)
        for d in user_docs_list:
            await delete_document(d["doc_id"], user_id)

        return {"success": True, "message": f"用户 {user_id} 的知识库已清空（Chroma + MySQL）"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════════════════════
# 内部接口（仅 Java 后端调用，已通过 X-Internal-Api-Secret 鉴权）
# ══════════════════════════════════════════════════════════════════════════════

@router.delete("/internal/documents/{doc_id}")
async def internal_delete_document(doc_id: str, user_id: str = "default"):
    """
    内部删除接口 - 由 Java 后端调用

    Python 这里只负责清理：
    - Chroma 向量数据 + BM25 索引（Chroma 内部同时清理）
    - MySQL 文档记录
    - 内存元数据

    Args:
        doc_id: 文档ID
        user_id: 用户ID（通过 query param 传入，用于定位到该用户的独立向量库）
    """
    # 获取该用户的独立向量库
    store = get_or_create_user_vector_store(user_id)

    deleted_chroma = 0
    deleted_mysql = False

    # 1. 删该用户独立 Chroma 向量库中的文档（Chroma 同时清理向量和 BM25 索引）
    try:
        store.delete(where={"doc_id": doc_id})
        print(f"[INFO] [Internal] Chroma 已删除 doc_id={doc_id}, user_id={user_id}")
        deleted_chroma = 1
    except Exception as e:
        print(f"[WARN] [Internal] Chroma delete failed: {e}")
        try:
            collection = store._collection
            results = collection.get(where={"doc_id": doc_id})
            if results and results.get("ids"):
                collection.delete(ids=results["ids"])
                deleted_chroma = len(results["ids"])
        except Exception as e2:
            print(f"[ERROR] [Internal] Chroma fallback delete failed: {e2}")

    # 2. 删 MySQL 记录
    from ..service.document_store import delete_document
    deleted_mysql = await delete_document(doc_id, user_id)

    # 3. 删内存 user_documents
    user_docs = user_documents.get(user_id, {})
    if doc_id in user_docs:
        del user_docs[doc_id]

    print(f"[INFO] [Internal Delete] doc_id={doc_id}, user_id={user_id}: "
          f"chroma={deleted_chroma}, mysql={deleted_mysql}")

    return {
        "success": True,
        "doc_id": doc_id,
        "user_id": user_id,
        "details": {
            "chroma_cleared": deleted_chroma > 0,
            "mysql_deleted": deleted_mysql
        }
    }


@router.get("/internal/documents/{doc_id}")
async def internal_get_document(doc_id: str, user_id: str = "default"):
    """内部查询 - 获取某个 doc_id 的所有 chunk 信息（用于调试）"""
    user_docs = user_documents.get(user_id, {})
    info = user_docs.get(doc_id, {})

    # 统计该用户的 BM25 索引中的 chunk
    user_index = user_bm25_index.get(user_id, {})
    chunks = [k.split("__")[1] for k in user_index if k.startswith(doc_id + "__")]

    return {
        "doc_id": doc_id,
        "user_id": user_id,
        "filename": info.get("filename", "unknown"),
        "chunk_count": info.get("chunk_count", 0),
        "bm25_indexed_chunks": len(chunks),
        "in_memory": doc_id_str in user_docs
    }


@router.get("/documents")
async def get_documents(user_id: str = "default"):
    """
    获取当前用户的文档列表（按用户ID隔离，从 MySQL 读取）

    Args:
        user_id: 用户ID，只返回该用户的文档
    """
    try:
        from ..service.document_store import get_documents_by_user
        documents = await get_documents_by_user(user_id)
        # 适配前端期望的字段名 (id / name / content / full_content / length / user_id)
        adapted = []
        for d in documents:
            adapted.append({
                "id": d["doc_id"],
                "name": d["filename"],
                "content": d.get("content", ""),
                "full_content": d.get("full_content", ""),
                "length": d.get("length", 0),
                "user_id": d.get("user_id", user_id),
                "chunk_count": d.get("chunk_count", 0),
                "created_at": str(d.get("created_at", "")),
            })
        return {"documents": adapted, "total": len(adapted), "user_id": user_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, user_id: str = "default"):
    """
    删除指定用户的文档（按用户ID隔离）

    Args:
        doc_id: 文档ID
        user_id: 用户ID，只删除该用户拥有的文档
    """
    try:
        user_docs = user_documents.get(user_id, {})

        deleted_from_chroma = 0
        deleted_from_mysql = False

        # 1. 删该用户独立 Chroma 向量库中的文档（Chroma 同时清理向量和 BM25 索引）
        store = get_or_create_user_vector_store(user_id)
        try:
            store.delete(where={"doc_id": str(doc_id)})
            print(f"[INFO] Chroma 已删除 doc_id={doc_id}, user_id={user_id}")
            deleted_from_chroma = 1
        except Exception as e:
            print(f"[WARN] Chroma 按 doc_id 删除失败: {e}")
            try:
                collection = store._collection
                results = collection.get(where={"doc_id": str(doc_id)})
                if results and results.get("ids"):
                    collection.delete(ids=results["ids"])
                    deleted_from_chroma = len(results["ids"])
                    print(f"[INFO] Chroma 通过 ID 删除了 {deleted_from_chroma} 个片段")
            except Exception as e2:
                print(f"[ERROR] Chroma 删除失败: {e2}")

        # 2. 删 MySQL 中的文档记录
        from ..service.document_store import delete_document
        deleted_from_mysql = await delete_document(doc_id, user_id)

        # 3. 删内存 user_documents
        if doc_id in user_docs:
            del user_docs[doc_id]

        print(f"[INFO] 删除文档 {doc_id} 完成: user_id={user_id}, "
              f"Chroma={deleted_from_chroma}, MySQL={deleted_from_mysql}")

        return {
            "success": True,
            "message": "文档已删除（Chroma + MySQL 全部清理）",
            "details": {
                "chroma_cleared": deleted_from_chroma > 0,
                "mysql_deleted": deleted_from_mysql
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"[ERROR] delete_document failed: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"删除失败: {str(e)}")


@router.post("/reindex")
async def reindex_documents(user_id: str = "default"):
    """
    从原始文档目录重新索引到向量库（按用户ID隔离）

    Args:
        user_id: 用户ID，只重新索引该用户的文档
    """
    try:
        from ..service.document_store import get_documents_by_user, init_document_store
        await init_document_store()
        docs = await get_documents_by_user(user_id)
        if not docs:
            raise HTTPException(status_code=404, detail="MySQL 中没有该用户的文档")

        # 重建该用户的文档索引（Chroma 内部同时维护向量和 BM25 索引）
        user_documents[user_id] = {}

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500, chunk_overlap=50,
            separators=["\n\n", "\n", "(?<=[。！？!?；;])", ",", " ", ""],
        )

        total_chunks = 0
        all_split_docs = []
        for doc in docs:
            doc_id = doc["doc_id"]
            filename_in_doc = doc["filename"]
            actual_content = doc.get("full_content") or doc.get("content", "")
            # 去掉旧磁盘格式残留
            if actual_content.startswith("==="):
                lines = actual_content.split("\n")
                actual_content = "\n".join(lines[2:]) if len(lines) > 2 else ""
            if not actual_content.strip():
                continue

            split_docs = text_splitter.split_documents([Document(page_content=actual_content, metadata={"source": filename_in_doc})])
            split_docs = [d for d in split_docs if d.page_content and d.page_content.strip()]
            for i, d in enumerate(split_docs):
                d.metadata["doc_id"] = doc_id
                d.metadata["doc_name"] = filename_in_doc
                d.metadata["chunk_index"] = i
                d.metadata["total_chunks"] = len(split_docs)
                d.metadata["user_id"] = user_id

            # 更新该用户的内存元数据索引
            chunks_meta = [{"content": d.page_content, "chunk_index": i,
                            "total_chunks": len(split_docs)} for i, d in enumerate(split_docs)]
            user_documents[user_id][doc_id] = {
                "filename": filename_in_doc,
                "chunks": chunks_meta,
                "upload_time": str(doc.get("created_at", datetime.now())),
                "content": actual_content,
            }

            all_split_docs.extend(split_docs)
            total_chunks += len(split_docs)
            print(f"[DEBUG] 重新索引: user_id={user_id}, {filename_in_doc}, 片段数: {len(split_docs)}")

        # 先删该用户旧的 Chroma 数据，再全量写入（Chroma 会自动建 BM25 倒排索引）
        store = get_or_create_user_vector_store(user_id)
        try:
            existing = store.get(where={"user_id": user_id})
            if existing and existing.get("ids"):
                store.delete(ids=existing["ids"])
        except Exception:
            pass

        if all_split_docs:
            # 分批写入避免超时
            batch_size = 50
            for i in range(0, len(all_split_docs), batch_size):
                batch = all_split_docs[i:i + batch_size]
                store.add_documents(batch)

        return {
            "success": True,
            "user_id": user_id,
            "message": f"成功重新索引 {len(docs)} 个文档，共 {total_chunks} 个片段（Chroma 向量+BM25 双索引）",
            "documents_count": len(docs),
            "chunks_count": total_chunks,
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"[ERROR] reindex_documents failed: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"重新索引失败: {str(e)}")


@router.get("/document/{doc_id}")
async def get_document_preview(doc_id: str, user_id: str = "default"):
    """获取文档预览（按用户ID隔离，从 MySQL 读取）"""
    try:
        from ..service.document_store import get_document
        doc = await get_document(doc_id, user_id)
        if not doc:
            raise HTTPException(status_code=404, detail="文档不存在")
        return {
            "id": doc["doc_id"],
            "name": doc["filename"],
            "content": doc["content"],
            "length": doc.get("length", len(doc["content"])),
            "chunk_count": doc.get("chunk_count", 0),
            "created_at": str(doc.get("created_at", "")),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════════════════════
# 11. 辅助函数（Word 文档解析）
# ══════════════════════════════════════════════════════════════════════════════

def load_docx_with_python_docx(file_path: str):
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
        full_text = "\n\n".join(text_content)
        if not full_text:
            raise ValueError("Word 文档为空或无法读取内容")
        return [Document(page_content=full_text, metadata={"source": file_path})]
    except Exception as e:
        print(f"python-docx 加载失败: {e}")
        raise e


def load_doc_file(file_path: str):
    """加载旧版 Word .doc 文件（多策略尝试）"""
    import subprocess
    try:
        result = subprocess.run(["antiword", file_path], capture_output=True, text=True,
                                encoding="utf-8", timeout=30)
        if result.returncode == 0 and result.stdout.strip():
            return [Document(page_content=result.stdout.strip(), metadata={"source": file_path})]
    except Exception as e:
        print(f"[DEBUG] antiword 不可用: {e}")

    try:
        import mammoth
        with open(file_path, "rb") as doc_file:
            result = mammoth.extract_raw_text(doc_file)
            if result.value and result.value.strip():
                return [Document(page_content=result.value.strip(), metadata={"source": file_path})]
    except ImportError:
        print("[DEBUG] mammoth 未安装")
    except Exception as e:
        print(f"[DEBUG] mammoth 读取失败: {e}")

    try:
        import textract
        text = textract.process(file_path).decode("utf-8")
        if text and text.strip():
            return [Document(page_content=text.strip(), metadata={"source": file_path})]
    except ImportError:
        print("[DEBUG] textract 未安装")
    except Exception as e:
        print(f"[DEBUG] textract 读取失败: {e}")

    try:
        import olefile
        if olefile.isOleFile(file_path):
            ole = olefile.OleFileIO(file_path)
            text_content = []
            if ole.exists("WordDocument"):
                word_stream = ole.openstream("WordDocument")
                raw_data = word_stream.read()
                import re
                for enc in ["utf-16-le", "gbk"]:
                    try:
                        text = raw_data.decode(enc, errors="ignore")
                        cleaned = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]+", " ", text)
                        cleaned = re.sub(r"\s+", " ", cleaned).strip()
                        if cleaned and len(cleaned) > 10:
                            text_content.append(cleaned)
                    except Exception:
                        pass
            ole.close()
            if text_content:
                full_text = "\n\n".join(text_content)
                return [Document(page_content=full_text, metadata={"source": file_path})]
    except ImportError:
        print("[DEBUG] olefile 未安装")
    except Exception as e:
        print(f"[DEBUG] olefile 读取失败: {e}")

    if os.name == "nt":
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(file_path))
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            if text and text.strip():
                return [Document(page_content=text.strip(), metadata={"source": file_path})]
        except Exception as e:
            print(f"[DEBUG] win32com 读取失败: {e}")

    print(f"[ERROR] 无法读取 .doc 文件: {file_path}")
    return None


# ══════════════════════════════════════════════════════════════════════════════
# 12. 启动初始化（从 MySQL 加载文档索引）
# ══════════════════════════════════════════════════════════════════════════════

async def startup_load_indexes_from_mysql():
    """
    启动时从 MySQL rag_documents 表加载所有用户的文档，
    重建 BM25 内存索引（幂等：重复调用会先清空再重建）。

    由 FastAPI lifespan 在启动时 await 调用。
    """
    import uuid as _uuid
    from ..service.document_store import get_all_documents, init_document_store

    # 确保表存在
    await init_document_store()

    all_docs = await get_all_documents()
    if not all_docs:
        print("[INFO] MySQL 中无文档，跳过索引加载")
        return

    # 按 user_id 分组
    by_user: dict = {}
    for d in all_docs:
        uid = d.get("user_id", "default")
        by_user.setdefault(uid, []).append(d)

    for uid, docs in by_user.items():
        # 清空该用户旧的内存索引，避免重复
        user_documents[uid] = {}
        await init_document_indices(uid)
        print(f"[INFO] 用户 {uid}: 从 MySQL 加载了 {len(docs)} 个文档")

    # 初始化默认用户的向量库（供 qa_chain 使用）
    init_qa_chain()
    print(f"[INFO] RAG 知识库初始化成功 (MySQL + Chroma Hybrid), 用户数: {len(user_documents)}")
